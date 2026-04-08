import csv
import io
import json
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse
from app.core.aws import get_report_from_s3, get_job
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()


def build_json(report: dict) -> bytes:
    return json.dumps(report, indent=2).encode("utf-8")


def build_csv(report: dict) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)

    # profile
    writer.writerow(["Section", "Field", "Value"])
    profile = report.get("profile", {})
    writer.writerow(["Profile", "Name", profile.get("name", "")])
    writer.writerow(["Profile", "Role", profile.get("current_role", "")])
    writer.writerow(["Profile", "Years Experience", profile.get("years_experience", "")])
    writer.writerow(["Profile", "Skills", ", ".join(profile.get("skills", []))])
    writer.writerow(["Profile", "Tools", ", ".join(profile.get("tools", []))])

    # risk
    risk = report.get("risk", {})
    writer.writerow(["Risk", "Score", risk.get("score", "")])
    writer.writerow(["Risk", "Band", risk.get("band", "")])
    writer.writerow(["Risk", "Highly Automatable %", risk.get("highly_automatable_pct", "")])
    writer.writerow(["Risk", "Partially Automatable %", risk.get("partially_automatable_pct", "")])
    writer.writerow(["Risk", "Low Automatable %", risk.get("low_automatable_pct", "")])
    writer.writerow(["Risk", "Human Critical %", risk.get("human_critical_pct", "")])

    # tasks
    writer.writerow([])
    writer.writerow(["Tasks", "Category", "Description"])
    for task in report.get("tasks", {}).get("tasks", []):
        writer.writerow(["Task", task.get("category", ""), task.get("description", "")])

    # recommendations
    recommendations = report.get("recommendations", {})
    for section, items in recommendations.items():
        writer.writerow([])
        writer.writerow([section.replace("_", " ").title(), "", ""])
        for item in items:
            writer.writerow(["", "", item])

    # roast
    roast = report.get("roast")
    if roast:
        writer.writerow([])
        writer.writerow(["Roast", "", roast])

    return output.getvalue().encode("utf-8")


def build_docx(report: dict) -> bytes:
    doc = Document()

    # styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def field(label: str, value: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    # title
    title = doc.add_heading("AutoRisk AI — Automation Risk Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # profile
    profile = report.get("profile", {})
    heading("Profile")
    field("Name", profile.get("name", ""))
    field("Role", profile.get("current_role", ""))
    field("Years of experience", str(profile.get("years_experience", "")))
    field("Skills", ", ".join(profile.get("skills", [])))
    field("Tools", ", ".join(profile.get("tools", [])))
    doc.add_paragraph()

    # risk score
    risk = report.get("risk", {})
    heading("Automation Risk Score")
    field("Score", f"{risk.get('score', '')} / 100")
    field("Band", risk.get("band", "").replace("_", " ").title())
    field("Highly automatable", f"{risk.get('highly_automatable_pct', '')}%")
    field("Partially automatable", f"{risk.get('partially_automatable_pct', '')}%")
    field("Low automatable", f"{risk.get('low_automatable_pct', '')}%")
    field("Human critical", f"{risk.get('human_critical_pct', '')}%")
    doc.add_paragraph()

    # tasks
    tasks = report.get("tasks", {})
    heading("Task Breakdown")
    for category, label in [
        ("highly_automatable", "Highly Automatable"),
        ("partially_automatable", "Partially Automatable"),
        ("low_automatable", "Low Automatable"),
        ("human_critical", "Human Critical"),
    ]:
        items = tasks.get(category, [])
        if items:
            heading(label, level=2)
            for item in items:
                bullet(item)
    doc.add_paragraph()

    # recommendations
    recommendations = report.get("recommendations", {})
    heading("Recommendations")
    sections = {
        "exposure_areas": "Exposure Areas",
        "resistant_strengths": "AI-Resistant Strengths",
        "upskill_roadmap": "Upskill Roadmap",
        "transition_paths": "Transition Paths",
    }
    for key, label in sections.items():
        items = recommendations.get(key, [])
        if items:
            heading(label, level=2)
            for item in items:
                bullet(item)
    doc.add_paragraph()

    # roast
    roast = report.get("roast")
    if roast:
        heading("Roast Mode")
        p = doc.add_paragraph(f'"{roast}"')
        p.runs[0].italic = True

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


@router.get("/download/{job_id}")
async def download_report(
    job_id: str,
    format: str = Query(default="json", enum=["json", "csv", "docx"]),
):
    # fetch job from DynamoDB
    job = await get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job.get("status") != "complete":
        raise HTTPException(status_code=400, detail="Report not ready yet.")

    s3_key = job.get("s3_key")
    if not s3_key:
        raise HTTPException(status_code=404, detail="Report file not found.")

    report = await get_report_from_s3(s3_key)

    if format == "json":
        return StreamingResponse(
            io.BytesIO(build_json(report)),
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=autorisk-{job_id[:8]}.json"},
        )

    elif format == "csv":
        return StreamingResponse(
            io.BytesIO(build_csv(report)),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=autorisk-{job_id[:8]}.csv"},
        )

    elif format == "docx":
        return StreamingResponse(
            io.BytesIO(build_docx(report)),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=autorisk-{job_id[:8]}.docx"},
        )